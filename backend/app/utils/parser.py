"""
Document Parser - Multi-format document extraction
Supports: PDF, DOCX, TXT, MD, Python/JS code files
"""
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field
from enum import Enum
import fitz  # PyMuPDF
from docx import Document
from loguru import logger


class DocumentType(str, Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MARKDOWN = "markdown"
    PYTHON = "python"
    JAVASCRIPT = "javascript"
    CODE = "code"
    UNKNOWN = "unknown"


@dataclass
class ParsedDocument:
    """Parsed document with extracted content."""
    filename: str
    doc_type: DocumentType
    content: str
    word_count: int
    metadata: dict = field(default_factory=dict)
    pages: list[str] = field(default_factory=list)
    
    @property
    def is_code(self) -> bool:
        """Check if document is code."""
        return self.doc_type in [
            DocumentType.PYTHON, 
            DocumentType.JAVASCRIPT, 
            DocumentType.CODE
        ]


class DocumentParser:
    """
    Multi-format document parser.
    Extracts text content from various file types.
    """
    
    EXTENSION_MAP = {
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".doc": DocumentType.DOCX,
        ".txt": DocumentType.TXT,
        ".md": DocumentType.MARKDOWN,
        ".py": DocumentType.PYTHON,
        ".js": DocumentType.JAVASCRIPT,
        ".ts": DocumentType.JAVASCRIPT,
        ".jsx": DocumentType.JAVASCRIPT,
        ".tsx": DocumentType.JAVASCRIPT,
        ".java": DocumentType.CODE,
        ".cpp": DocumentType.CODE,
        ".c": DocumentType.CODE,
        ".go": DocumentType.CODE,
        ".rs": DocumentType.CODE,
    }
    
    def __init__(self):
        self.parsers = {
            DocumentType.PDF: self._parse_pdf,
            DocumentType.DOCX: self._parse_docx,
            DocumentType.TXT: self._parse_text,
            DocumentType.MARKDOWN: self._parse_text,
            DocumentType.PYTHON: self._parse_code,
            DocumentType.JAVASCRIPT: self._parse_code,
            DocumentType.CODE: self._parse_code,
        }
    
    def parse(self, file_path: str | Path) -> ParsedDocument:
        """
        Parse a document and extract its content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParsedDocument with extracted content
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        doc_type = self._detect_type(path)
        logger.info(f"Parsing {path.name} as {doc_type.value}")
        
        parser = self.parsers.get(doc_type, self._parse_text)
        return parser(path, doc_type)
    
    def parse_content(self, content: str, filename: str = "content.txt") -> ParsedDocument:
        """
        Parse raw text content (for code snippets, etc).
        
        Args:
            content: Raw text content
            filename: Virtual filename for type detection
            
        Returns:
            ParsedDocument with the content
        """
        path = Path(filename)
        doc_type = self._detect_type(path)
        
        return ParsedDocument(
            filename=filename,
            doc_type=doc_type,
            content=content,
            word_count=len(content.split()),
            metadata={"source": "direct_content"},
            pages=[content]
        )
    
    def _detect_type(self, path: Path) -> DocumentType:
        """Detect document type from file extension."""
        suffix = path.suffix.lower()
        return self.EXTENSION_MAP.get(suffix, DocumentType.UNKNOWN)
    
    def _parse_pdf(self, path: Path, doc_type: DocumentType) -> ParsedDocument:
        """Parse PDF documents using PyMuPDF."""
        doc = fitz.open(str(path))
        pages = []
        full_text = []
        
        for page in doc:
            text = page.get_text("text")
            pages.append(text)
            full_text.append(text)
        
        content = "\n\n".join(full_text)
        
        metadata = {
            "page_count": len(pages),
            "author": doc.metadata.get("author", ""),
            "title": doc.metadata.get("title", "") or path.stem,
            "creation_date": doc.metadata.get("creationDate", ""),
        }
        
        doc.close()
        
        return ParsedDocument(
            filename=path.name,
            doc_type=doc_type,
            content=content,
            word_count=len(content.split()),
            metadata=metadata,
            pages=pages
        )
    
    def _parse_docx(self, path: Path, doc_type: DocumentType) -> ParsedDocument:
        """Parse Word documents."""
        doc = Document(str(path))
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        content = "\n\n".join(paragraphs)
        
        # Also extract tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        
        if tables_text:
            content += "\n\n[Tables]\n" + "\n".join(tables_text)
        
        return ParsedDocument(
            filename=path.name,
            doc_type=doc_type,
            content=content,
            word_count=len(content.split()),
            metadata={"paragraph_count": len(paragraphs), "table_count": len(doc.tables)},
            pages=[content]
        )
    
    def _parse_text(self, path: Path, doc_type: DocumentType) -> ParsedDocument:
        """Parse plain text and markdown files."""
        content = path.read_text(encoding="utf-8")
        
        return ParsedDocument(
            filename=path.name,
            doc_type=doc_type,
            content=content,
            word_count=len(content.split()),
            metadata={},
            pages=[content]
        )
    
    def _parse_code(self, path: Path, doc_type: DocumentType) -> ParsedDocument:
        """Parse code files with syntax-aware metadata."""
        content = path.read_text(encoding="utf-8")
        
        # Extract basic code metrics
        lines = content.split("\n")
        code_lines = [l for l in lines if l.strip() and not l.strip().startswith(("#", "//", "/*", "*", "'''", '"""'))]
        
        metadata = {
            "total_lines": len(lines),
            "code_lines": len(code_lines),
            "language": doc_type.value,
            "has_imports": any("import" in l or "require" in l or "from" in l for l in lines[:20])
        }
        
        # For Python, try to extract function/class names
        if doc_type == DocumentType.PYTHON:
            functions = [l.strip() for l in lines if l.strip().startswith("def ")]
            classes = [l.strip() for l in lines if l.strip().startswith("class ")]
            metadata["functions"] = len(functions)
            metadata["classes"] = len(classes)
        
        return ParsedDocument(
            filename=path.name,
            doc_type=doc_type,
            content=content,
            word_count=len(content.split()),
            metadata=metadata,
            pages=[content]
        )


# Singleton instance
parser = DocumentParser()
